from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create a new Document
doc = Document()

# Set document title
doc.add_heading('[Full Name]', level=1)

# Contact Information
doc.add_paragraph('Contact: 971582360459\nWhatsApp: 9496808698\nEmail: ansabaansi1234@gmail.com\nDate of Birth: 12/10/2000\nMarital Status: Married\nCitizenship: Indian\nPassport Number: S7864844')

# Add section for Professional Summary
doc.add_heading('Professional Summary', level=2)
doc.add_paragraph(
    "Motivated and dedicated psychology graduate with a diploma in counseling psychology, specializing in NLP (Neuro-Linguistic Programming) and hypnotherapy. With experience in counseling, teaching, and administrative roles, I aim to contribute my communication, leadership, and problem-solving skills to an organization where I can actively participate in its growth and development."
)

# Add Skills section
doc.add_heading('Skills', level=2)

# Soft Skills
doc.add_heading('Soft Skills', level=3)
doc.add_paragraph(
    "• Strong communication skills\n"
    "• Effective teamwork and collaboration\n"
    "• High adaptability to new environments and challenges\n"
    "• Leadership experience in training and counseling\n"
    "• Problem-solving expertise\n"
    "• Customer service oriented\n"
    "• Dependable and hardworking with a flexible, positive attitude"
)

# Technical Skills
doc.add_heading('Technical Skills', level=3)
doc.add_paragraph(
    "• Knowledge of administrative procedures\n"
    "• Skilled in drafting communication letters\n"
    "• Proficient in Microsoft Office Suite\n"
    "• Trained in NLP and hypnotherapy techniques\n"
    "• Experienced in counseling psychology and public relations"
)

# Experience Section
doc.add_heading('Experience', level=2)

# NLP Trainer & Clinical Hypnotist
doc.add_heading('NLP Trainer & Clinical Hypnotist', level=3)
doc.add_paragraph('Independent (Jan 2023 – Jun 2023)')
doc.add_paragraph(
    "• Provided clinical hypnotherapy services and NLP-based counseling\n"
    "• Trained individuals in NLP techniques for personal development\n"
    "• Assisted clients in achieving their personal and professional goals through tailored therapeutic interventions"
)

# Counselor
doc.add_heading('Counselor', level=3)
doc.add_paragraph('Ease Edu Hub (Jan 2023 – Jun 2023)')
doc.add_paragraph(
    "• Offered psychological support and counseling services to students\n"
    "• Worked closely with clients to address academic and emotional concerns\n"
    "• Managed individual counseling sessions for students and parents"
)

# Office Staff
doc.add_heading('Office Staff', level=3)
doc.add_paragraph('Myntra Pilathara (Aug 2023 – Nov 2023)')
doc.add_paragraph(
    "• Managed administrative tasks including documentation, filing, and data entry\n"
    "• Coordinated internal communication and supported office management\n"
    "• Assisted in customer service operations and office workflow optimization"
)

# Education Section
doc.add_heading('Education', level=2)
doc.add_paragraph(
    "• 2022: Diploma in Counseling Psychology (Specialization in NLP and Hypnotherapy)\n"
    "• 2021: Bachelor of Arts in Psychology (Indira Gandhi National Open University - IGNOU)\n"
    "• 2018: Plus Two (Higher Secondary Education)\n"
    "• 2016: Secondary School Leaving Certificate (SSLC)"
)

# Areas of Interest Section
doc.add_heading('Areas of Interest', level=2)
doc.add_paragraph("• Counseling\n• Teaching\n• Public Relations")

# Save the document
doc.save('CV.docx')

print("CV has been created and saved as 'CV.docx'.")